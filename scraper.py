import os, pathlib, hashlib, re, collections
from datetime import datetime
from bs4 import BeautifulSoup
from playwright.async_api import async_playwright, TimeoutError
import pandas as pd

# ------------------------- CONSTANTS -------------------------
BASE = "https://bis.gov.lv"
LIST_URL = BASE + "/bisp/lv/planned_constructions/list?page={page}"

# Pages to fetch (1..N). You can override in workflow env (PAGES_TOTAL)
PAGES_TOTAL = int(os.getenv("PAGES_TOTAL", "300"))

DEBUG_DIR = pathlib.Path("debug"); DEBUG_DIR.mkdir(parents=True, exist_ok=True)
REPORTS = pathlib.Path("reports"); REPORTS.mkdir(parents=True, exist_ok=True)

# Your filters (we normalize values before comparing)
AUTHORITIES_WHITELIST = {
    "RĪGAS VALSTSPILSĒTAS PAŠVALDĪBAS PILSĒTAS ATTĪSTĪBAS DEPARTAMENTS",
    "Ādažu novada būvvalde",
    "Saulkrastu novada būvvalde",
    "Ropažu novada pašvaldības būvvalde",
    "Siguldas novada būvvalde",
    "Salaspils novada pašvaldības iestāde \"Salaspils novada Būvvalde\"",
    "Ogres novada pašvaldības centrālās administrācijas Ogres novada būvvalde",
    "Ķekavas novada pašvaldības būvvalde",
    "OLAINES NOVADA PAŠVALDĪBAS BŪVVALDE",
    "Mārupes novada Būvvalde",
    "Jūrmalas Būvvalde",
}

PHASE_KEEP = {
    "Iecere",
    "Būvniecības ieceres publiskā apspriešana",
    "Projektēšanas nosacījumu izpilde",
    "Būvdarbu uzsākšanas nosacījumu izpilde",
}

TYPE_KEEP = {
    "Atjaunošana",
    "Vienkāršota atjaunošana",
    "Jauna būvniecība",
    "Pārbūve",
    "Vienkāršota pārbūve",
}

# Header labels used in the /list?page=N blocks
HEADER_MAP = {
    "Būvniecības kontroles institūcija": "authority",
    "Lietas numurs": "bis_number",
    "Būves nosaukums": "object",
    "Adrese": "address",
    "Būvniecības veids": "construction_type",
    "Būvniecības lietas stadija": "phase",
}
# -------------------------------------------------------------


# ------------------------- NORMALIZATION -------------------------
NBSP = "\u00A0"

def norm(s: str) -> str:
    """Normalize text for robust equality: replace NBSP, collapse spaces, strip."""
    if s is None:
        return ""
    s = s.replace(NBSP, " ")
    s = re.sub(r"\s+", " ", s)
    return s.strip()

# Pre-normalize filters once
AUTHORITIES_NORM = {norm(a): a for a in AUTHORITIES_WHITELIST}
PHASE_KEEP_NORM   = {norm(x) for x in PHASE_KEEP}
TYPE_KEEP_NORM    = {norm(x) for x in TYPE_KEEP}
# ---------------------------------------------------------------


# ------------------------- PARSING -------------------------
def stable_row_id(r: dict) -> str:
    """Prefer BIS number; otherwise hash a few stable fields."""
    if r.get("bis_number"):
        return f"bis:{r['bis_number']}"
    key = "|".join(str(r.get(k, "")) for k in ["authority","address","object","phase","construction_type"])
    return "h:" + hashlib.sha256(key.encode("utf-8")).hexdigest()[:16]

def extract_value(cell, header_text: str) -> str:
    """
    Values in .flextable__value also include a screen-reader label like 'Label: Value'.
    Strip 'Label:' prefix so we keep only the value.
    """
    val_el = cell.select_one(".flextable__value")
    t = norm(val_el.get_text(" ", strip=True) if val_el else "")
    prefix = header_text + ":"
    if t.startswith(prefix):
        t = norm(t[len(prefix):])
    return t

def parse_page(html: str):
    """
    Return (total_rows_on_page, matched_rows_list, diag Counter by authority).
    - total_rows_on_page: count of .flextable__row on the page (even if later filtered out)
    - matched_rows_list: rows that pass all filters
    - diag: diagnostic counts by authority (normalized), for visibility in logs
    """
    soup = BeautifulSoup(html, "lxml")
    row_nodes = soup.select(".flextable__row")
    total_rows = len(row_nodes)
    out = []
    diag = collections.Counter()

    for row in row_nodes:
        rec = {"details_url": None}
        for cell in row.select(".flextable__cell"):
            header = norm(cell.get("data-column-header-name") or "")
            key = HEADER_MAP.get(header)
            if not key:
                continue

            text = extract_value(cell, header)
            a = cell.select_one("a.public_list__link[href]")
            if key == "bis_number" and a:
                href = a.get("href", "")
                if href.startswith("/"):
                    href = BASE + href
                rec["details_url"] = href
                text = norm(a.get_text(" ", strip=True))
            rec[key] = text

        auth_n = norm(rec.get("authority"))
        if auth_n:
            diag[auth_n] += 1

        # Apply normalized filters
        if auth_n not in AUTHORITIES_NORM:
            continue
        phase_n = norm(rec.get("phase"))
        if phase_n and phase_n not in PHASE_KEEP_NORM:
            continue
        type_n = norm(rec.get("construction_type"))
        if type_n and type_n not in TYPE_KEEP_NORM:
            continue

        # Keep canonical authority label & normalized values
        rec["authority"] = AUTHORITIES_NORM.get(auth_n, rec.get("authority"))
        rec["phase"] = phase_n
        rec["construction_type"] = type_n

        rec["id"] = stable_row_id(rec)
        out.append(rec)

    return total_rows, out, diag
# -----------------------------------------------------------


# ------------------------- REPORTS -------------------------
def make_html_report(rows: list[dict], pages_seen: int, scanned: int) -> str:
    # Build a clean HTML table with clickable BIS numbers
    if not rows:
        body = "<p>Nav ierakstu, kas atbilst filtriem.</p>"
    else:
        df = pd.DataFrame(rows)[[
            "bis_number", "authority", "address", "object", "phase", "construction_type", "details_url"
        ]].copy()
        df["BIS lieta"] = df.apply(
            lambda r: (f'<a href="{r["details_url"]}" target="_blank" rel="noopener">{r["bis_number"]}</a>'
                       if r.get("details_url") and r.get("bis_number") else (r.get("bis_number",""))),
            axis=1
        )
        df.rename(columns={
            "authority": "Būvniecības kontroles institūcija",
            "address": "Adrese",
            "object": "Būves nosaukums",
            "phase": "Būvniecības lietas stadija",
            "construction_type": "Būvniecības veids",
        }, inplace=True)
        df = df[["BIS lieta", "Būvniecības kontroles institūcija", "Adrese", "Būves nosaukums",
                 "Būvniecības lietas stadija", "Būvniecības veids"]]
        body = df.to_html(index=False, escape=False)

    meta = f"""
    <p><strong>Pārlapotas lapas:</strong> {pages_seen} &nbsp;|&nbsp;
       <strong>Rindas skenētas kopā:</strong> {scanned} &nbsp;|&nbsp;
       <strong>Atlasīto ierakstu skaits:</strong> {len(rows)}</p>
    """
    css = """
    <style>
      body{font-family:system-ui,-apple-system,Segoe UI,Roboto,Arial,sans-serif; padding:16px;}
      table{border-collapse:collapse; width:100%; font-size:14px;}
      th,td{border:1px solid #e5e7eb; padding:8px; vertical-align:top;}
      th{background:#f3f4f6; text-align:left;}
      a{ text-decoration:none; }
    </style>
    """
    return f"""<!doctype html><meta charset="utf-8"><title>BIS atskaite</title>{css}
    <h1>BIS plānoto būvniecību atskaite</h1>
    <p><small>{datetime.now().strftime('%Y-%m-%d %H:%M')}</small></p>
    {meta}
    {body}
    """

def make_docx_report(rows: list[dict], outfile: pathlib.Path):
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)
    doc.add_heading('BIS plānoto būvniecību atskaite', level=1)

    if not rows:
        doc.add_paragraph("Nav ierakstu, kas atbilst filtriem.")
        doc.save(str(outfile))
        return

    cols = ["BIS lieta", "Būvniecības kontroles institūcija", "Adrese",
            "Būves nosaukums", "Būvniecības lietas stadija", "Būvniecības veids"]
    table = doc.add_table(rows=1, cols=len(cols))
    hdr = table.rows[0].cells
    for i,c in enumerate(cols): hdr[i].text = c

    def add_hyperlink(cell_paragraph, url, text):
        part = cell_paragraph.part
        r_id = part.relate_to(url,
                              reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                              is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(_qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(_qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)
        new_run.append(rPr)
        t = OxmlElement('w:t'); t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        cell_paragraph._p.append(hyperlink)

    for r in rows:
        row = table.add_row().cells
        # BIS lieta with link
        cell0 = row[0].paragraphs[0]
        if r.get("details_url") and r.get("bis_number"):
            add_hyperlink(cell0, r["details_url"], r["bis_number"])
        else:
            cell0.add_run(r.get("bis_number",""))
        row[1].text = r.get("authority","")
        row[2].text = r.get("address","")
        row[3].text = r.get("object","")
        row[4].text = r.get("phase","")
        row[5].text = r.get("construction_type","")

    doc.save(str(outfile))

def save_reports(rows: list[dict], pages_seen: int, scanned: int, diag_accum: collections.Counter):
    # CSV (Excel-friendly, Latvian-safe)
    df = pd.DataFrame(rows)
    if not df.empty:
        df.drop_duplicates(subset=["id"], inplace=True)
    df.to_csv(REPORTS / "latest.csv", index=False, encoding="utf-8-sig")
    today = datetime.now().strftime("%Y-%m-%d")
    df.to_csv(REPORTS / f"{today}.csv", index=False, encoding="utf-8-sig")

    # CHANGELOG (simple)
    (REPORTS / "CHANGELOG.md").write_text(
        "# Snapshot {}\n\n- Pārlapotas lapas: {}\n- Rindas skenētas kopā: {}\n- Rindas pēc filtriem (unikālas): {}\n".format(
        datetime.now().strftime("%Y-%m-%d %H:%M"), pages_seen, scanned, 0 if df.empty else len(df)),
        encoding="utf-8"
    )

    # HTML
    html = make_html_report(rows, pages_seen, scanned)
    (REPORTS / "report.html").write_text(html, encoding="utf-8")

    # DOCX (best-effort)
    try:
        make_docx_report(rows, REPORTS / "report.docx")
    except Exception as e:
        (REPORTS / "report.docx.error.txt").write_text(str(e), encoding="utf-8")
# -----------------------------------------------------------


# ------------------------- MAIN -------------------------
async def main():
    async with async_playwright() as p:
        browser = await p.chromium.launch()
        ctx = await browser.new_context()
        page = await ctx.new_page()

        all_rows = []
        total_scanned = 0
        pages_fetched = 0
        diag_all = collections.Counter()

        for n in range(1, PAGES_TOTAL + 1):
            url = LIST_URL.format(page=n)
            try:
                await page.goto(url, wait_until="domcontentloaded", timeout=180000)
            except TimeoutError:
                await page.goto(url, wait_until="domcontentloaded", timeout=180000)

            # Accept cookies if shown
            try:
                for t in ["Apstiprināt", "Apstiprināt izvēlētās", "Apstiprināt visas", "Piekrītu"]:
                    btn = page.get_by_text(t, exact=False).first
                    if await btn.count() > 0:
                        await btn.click(timeout=1200)
                        break
            except:
                pass

            await page.wait_for_timeout(250)
            html = await page.content()

            if n <= 2:
                (DEBUG_DIR / f"page-{n}.html").write_text(html, encoding="utf-8")

            total_rows, matched, diag = parse_page(html)
            pages_fetched += 1
            total_scanned += total_rows
            diag_all.update(diag)

            # Only stop at true end-of-list (no rows at all)
            if total_rows == 0:
                break

            all_rows.extend(matched)

        await browser.close()

    save_reports(all_rows, pages_fetched, total_scanned, diag_all)
    print({
        "pages": pages_fetched,
        "rows_scanned": total_scanned,
        "rows_matched": len(all_rows),
        "top_authorities_seen": diag_all.most_common(5),
    })

if __name__ == "__main__":
    import asyncio
    asyncio.run(main())
